#!/usr/bin/env python3
"""
テスト用ファイル生成スクリプト
- Excel（複数シート、「テストシート」を含む）
- Word
- TIFF（縦長・横長の2種類）
"""

import os

def create_test_excel():
    """テスト用Excelファイル作成"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    wb = Workbook()
    
    # シート1: ダミーシート
    ws1 = wb.active
    ws1.title = "ダミーシート1"
    ws1['A1'] = "これはダミーシートです"
    ws1['A2'] = "このシートは変換対象外です"
    
    # シート2: テストシート（変換対象）
    ws2 = wb.create_sheet("テストシート")
    
    # ヘッダースタイル
    header_fill = PatternFill('solid', fgColor='4472C4')
    header_font = Font(bold=True, color='FFFFFF', size=12)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # タイトル
    ws2.merge_cells('A1:E1')
    ws2['A1'] = '設計変更通知一覧（テストデータ）'
    ws2['A1'].font = Font(bold=True, size=16)
    ws2['A1'].alignment = Alignment(horizontal='center')
    
    # ヘッダー行
    headers = ['No', '設変番号', '設変名称', 'ステータス', '担当者']
    for col, header in enumerate(headers, start=1):
        cell = ws2.cell(row=3, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')
    
    # テストデータ
    data = [
        (1, 'SH-2025-001', '材料変更（アルミ→ステンレス）', '承認済', '田中太郎'),
        (2, 'SH-2025-002', '寸法変更（φ10→φ12）', '審査中', '鈴木花子'),
        (3, 'SH-2025-003', '塗装色変更（白→シルバー）', '受付済', '佐藤次郎'),
        (4, 'SH-2025-004', 'サプライヤー変更', '保留', '高橋三郎'),
        (5, 'SH-2025-005', '製造工程変更', '完了', '伊藤四郎'),
    ]
    
    for row_idx, row_data in enumerate(data, start=4):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center')
    
    # 列幅調整
    ws2.column_dimensions['A'].width = 5
    ws2.column_dimensions['B'].width = 15
    ws2.column_dimensions['C'].width = 35
    ws2.column_dimensions['D'].width = 12
    ws2.column_dimensions['E'].width = 15
    
    # シート3: もう一つのダミー
    ws3 = wb.create_sheet("ダミーシート2")
    ws3['A1'] = "これもダミーシートです"
    
    # 保存
    output_path = 'input/test_excel.xlsx'
    wb.save(output_path)
    print(f'✓ Excel作成完了: {output_path}')
    print(f'  シート: {wb.sheetnames}')

def create_test_word():
    """テスト用Wordファイル作成"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # タイトル
    title = doc.add_heading('PDF変換テストドキュメント', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 概要
    doc.add_heading('1. 概要', level=1)
    doc.add_paragraph(
        'このドキュメントはPDF変換機能のテスト用に作成されました。'
        'LibreOfficeのheadlessモードを使用して、'
        'Wordファイル(.docx)からPDFへの変換を検証します。'
    )
    
    # テスト項目
    doc.add_heading('2. テスト項目', level=1)
    items = [
        'テキストの変換',
        '段落フォーマットの保持',
        '表の変換',
        '日本語文字の表示',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 表
    doc.add_heading('3. サンプル表', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # ヘッダー行
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '項目'
    hdr_cells[1].text = '値'
    hdr_cells[2].text = '備考'
    
    # データ行
    data = [
        ('変換形式', 'DOCX → PDF', 'LibreOffice使用'),
        ('文字コード', 'UTF-8', '日本語対応'),
        ('ステータス', 'テスト中', '検証用'),
    ]
    for i, row_data in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    # フッター
    doc.add_paragraph()
    footer = doc.add_paragraph('作成日: 2025年1月29日')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存
    output_path = 'input/test_word.docx'
    doc.save(output_path)
    print(f'✓ Word作成完了: {output_path}')

def create_test_tiffs():
    """テスト用TIFFファイル作成（縦長・横長）"""
    from PIL import Image, ImageDraw, ImageFont
    
    # フォント（システムフォントを使用、なければデフォルト）
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 40)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except:
        font = ImageFont.load_default()
        font_small = font
    
    # 1. 横長TIFF（幅 > 高さ）→ 回転不要
    img_landscape = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img_landscape)
    draw.rectangle([20, 20, 780, 580], outline='blue', width=3)
    draw.text((250, 250), "Landscape Image", fill='blue', font=font)
    draw.text((280, 320), "(800 x 600 - No Rotation)", fill='gray', font=font_small)
    draw.text((50, 500), "Width > Height", fill='darkblue', font=font_small)
    
    output_landscape = 'input/test_landscape.tiff'
    img_landscape.save(output_landscape, format='TIFF')
    print(f'✓ TIFF(横長)作成完了: {output_landscape} (800x600)')
    
    # 2. 縦長TIFF（高さ > 幅）→ 右90度回転が必要
    img_portrait = Image.new('RGB', (600, 800), color='lightyellow')
    draw = ImageDraw.Draw(img_portrait)
    draw.rectangle([20, 20, 580, 780], outline='red', width=3)
    draw.text((150, 350), "Portrait Image", fill='red', font=font)
    draw.text((130, 420), "(600 x 800 - Needs Rotation)", fill='gray', font=font_small)
    draw.text((50, 700), "Height > Width -> Rotate 90 deg", fill='darkred', font=font_small)
    
    output_portrait = 'input/test_portrait.tiff'
    img_portrait.save(output_portrait, format='TIFF')
    print(f'✓ TIFF(縦長)作成完了: {output_portrait} (600x800)')

def main():
    # 入力ディレクトリ作成
    os.makedirs('input', exist_ok=True)
    os.makedirs('output', exist_ok=True)
    os.makedirs('temp', exist_ok=True)
    
    print('=' * 50)
    print('テストファイル生成開始')
    print('=' * 50)
    
    # Excel作成
    print('\n[1/3] Excel作成中...')
    create_test_excel()
    
    # Word作成
    print('\n[2/3] Word作成中...')
    create_test_word()
    
    # TIFF作成
    print('\n[3/3] TIFF作成中...')
    create_test_tiffs()
    
    print('\n' + '=' * 50)
    print('テストファイル生成完了')
    print('=' * 50)
    print('\n生成されたファイル:')
    for f in os.listdir('input'):
        path = os.path.join('input', f)
        size_kb = os.path.getsize(path) / 1024
        print(f'  - input/{f} ({size_kb:.1f} KB)')

if __name__ == '__main__':
    main()
